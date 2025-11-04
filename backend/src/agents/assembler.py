"""Assembler agent for DOCX export.

Merges generated sections into a complete proposal document with citations.
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import List, Dict, Any, Optional
from datetime import datetime
import logging
import re

logger = logging.getLogger(__name__)


class Assembler:
    """Assembles proposal sections into formatted DOCX document"""
    
    def __init__(self):
        """Initialize assembler"""
        pass
    
    def assemble_proposal(
        self,
        sections: List[Dict[str, Any]],
        funding_call_name: Optional[str] = None,
        program_name: Optional[str] = None
    ) -> Document:
        """Assemble sections into a DOCX document.
        
        Args:
            sections: List of section data dicts with keys:
                - section_name: Section heading
                - text: Section content with inline citations
                - word_count: Number of words
                - citations: List of citation metadata
            funding_call_name: Name of funding call (for title page)
            program_name: Program/organization name (for title page)
            
        Returns:
            python-docx Document object
        """
        logger.info(f"[ASSEMBLER] ========== Starting document assembly ==========")
        logger.info(f"[ASSEMBLER] Sections to assemble: {len(sections)}")
        logger.info(f"[ASSEMBLER] Funding call: {funding_call_name}")
        logger.info(f"[ASSEMBLER] Program name: {program_name}")
        
        # Log section details
        for i, section in enumerate(sections):
            logger.info(
                f"[ASSEMBLER] Section {i+1}: '{section.get('section_name', 'Unknown')}' "
                f"({section.get('word_count', 0)} words, "
                f"{len(section.get('citations', []))} citations)"
            )
        
        try:
            # Create new document
            logger.debug("[ASSEMBLER] Creating new Document object")
            doc = Document()
            
            # Add title page
            logger.debug("[ASSEMBLER] Adding title page")
            self._add_title_page(
                doc,
                funding_call_name=funding_call_name,
                program_name=program_name
            )
            
            # Add page break after title
            logger.debug("[ASSEMBLER] Adding page break after title")
            doc.add_page_break()
            
            # Add each section
            for i, section_data in enumerate(sections):
                logger.info(
                    f"[ASSEMBLER] Processing section {i+1}/{len(sections)}: "
                    f"'{section_data.get('section_name', 'Unknown')}'"
                )
                
                try:
                    self._add_section(doc, section_data)
                    logger.debug(f"[ASSEMBLER] Section {i+1} added successfully")
                except Exception as e:
                    logger.error(
                        f"[ASSEMBLER] Error adding section {i+1} "
                        f"'{section_data.get('section_name', 'Unknown')}': {e}",
                        exc_info=True
                    )
                    raise
                
                # Add spacing between sections (but not after last section)
                if i < len(sections) - 1:
                    doc.add_paragraph()  # Blank line between sections
                    logger.debug(f"[ASSEMBLER] Added spacing after section {i+1}")
            
            logger.info(f"[ASSEMBLER] ========== Document assembly complete ==========")
            logger.info(f"[ASSEMBLER] Total paragraphs in document: {len(doc.paragraphs)}")
            return doc
            
        except Exception as e:
            logger.error(f"[ASSEMBLER] Document assembly failed: {e}", exc_info=True)
            raise
    
    def _add_title_page(
        self,
        doc: Document,
        funding_call_name: Optional[str] = None,
        program_name: Optional[str] = None
    ):
        """Add title page to document.
        
        Args:
            doc: Document object
            funding_call_name: Funding call title
            program_name: Program/organization name
        """
        # Title
        title = doc.add_heading(
            funding_call_name or "Grant Proposal",
            level=0
        )
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add spacing
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Program name (if provided)
        if program_name:
            program_para = doc.add_paragraph(program_name)
            program_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            program_run = program_para.runs[0]
            program_run.font.size = Pt(14)
            doc.add_paragraph()
        
        # Generation date
        date_para = doc.add_paragraph(
            f"Generated: {datetime.utcnow().strftime('%B %d, %Y')}"
        )
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_run = date_para.runs[0]
        date_run.font.size = Pt(12)
        
        # Footer note
        doc.add_paragraph()
        doc.add_paragraph()
        footer_para = doc.add_paragraph(
            "This document was generated using EasyGrant Smart Proposal Assistant. "
            "All citations reference uploaded supporting documents."
        )
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer_para.runs[0]
        footer_run.font.size = Pt(10)
        footer_run.italic = True
    
    def _add_section(self, doc: Document, section_data: Dict[str, Any]):
        """Add a section to the document.
        
        Args:
            doc: Document object
            section_data: Section metadata and content
        """
        section_name = section_data.get('section_name', 'Untitled Section')
        text = section_data.get('text', '')
        word_count = section_data.get('word_count', 0)
        
        logger.debug(f"[ASSEMBLER] _add_section called for '{section_name}'")
        logger.debug(f"[ASSEMBLER] Section text length: {len(text)} chars")
        logger.debug(f"[ASSEMBLER] Word count: {word_count}")
        
        # Add section heading
        logger.debug(f"[ASSEMBLER] Adding heading for '{section_name}'")
        heading = doc.add_heading(section_name, level=1)
        heading_run = heading.runs[0]
        heading_run.font.size = Pt(16)
        
        # Add word count indicator (optional, can be removed if not desired)
        if word_count > 0:
            logger.debug(f"[ASSEMBLER] Adding word count indicator: {word_count}")
            count_para = doc.add_paragraph(f"Word count: {word_count}")
            count_run = count_para.runs[0]
            count_run.font.size = Pt(9)
            count_run.italic = True
            count_run.font.color.rgb = None  # Use default color
        
        # Add section content with inline citations preserved
        logger.debug(f"[ASSEMBLER] Adding formatted text for '{section_name}'")
        self._add_formatted_text(doc, text)
        logger.debug(f"[ASSEMBLER] Section '{section_name}' completed")
    
    def _add_formatted_text(self, doc: Document, text: str):
        """Add text to document, preserving paragraph structure.
        
        Inline citations in format [Document, p.N] are preserved as-is.
        
        Args:
            doc: Document object
            text: Text content with inline citations
        """
        logger.debug(f"[ASSEMBLER] Adding formatted text with {len(text)} characters")
        
        # Split into paragraphs (double newline separator)
        paragraphs = text.split('\n\n')
        logger.debug(f"[ASSEMBLER] Split into {len(paragraphs)} paragraphs")
        
        for idx, para_text in enumerate(paragraphs):
            para_text = para_text.strip()
            if not para_text:
                logger.debug(f"[ASSEMBLER] Skipping empty paragraph at index {idx}")
                continue
            
            # Check if this is a heading line (starts with ##, ###, etc.)
            if para_text.startswith('#'):
                # Extract heading level and text
                heading_match = re.match(r'^(#{1,6})\s+(.+)$', para_text)
                if heading_match:
                    level = len(heading_match.group(1))
                    heading_text = heading_match.group(2)
                    logger.debug(f"[ASSEMBLER] Adding heading level {level}: {heading_text}")
                    # Add heading directly (don't create paragraph first)
                    doc.add_heading(heading_text, level=level)
                    continue
            
            # Add paragraph with text and citations
            logger.debug(f"[ASSEMBLER] Adding paragraph {idx+1}: {para_text[:50]}...")
            para = doc.add_paragraph()
            self._add_text_with_citation_highlighting(para, para_text)
    
    def _add_text_with_citation_highlighting(self, para, text: str):
        """Add text to paragraph, removing inline citations.
        
        Inline citations are stripped from the exported document as they are
        for internal verification only.
        
        Args:
            para: Paragraph object
            text: Text with inline citations
        """
        # Pattern to match inline citations: [Document, p.N] or [Document, p. N]
        # Allows optional space after 'p.' to match both formats
        citation_pattern = r'\[([^\]]+?),\s*p\.\s*(\d+)\]'
        
        citations_found = list(re.finditer(citation_pattern, text))
        logger.debug(f"[ASSEMBLER] Found {len(citations_found)} citations to remove from text")
        
        # Remove all inline citations from text
        clean_text = re.sub(citation_pattern, '', text)
        
        # Clean up any double spaces left after removing citations
        clean_text = re.sub(r'\s+', ' ', clean_text).strip()
        
        # Add the clean text without citations
        para.add_run(clean_text)
        logger.debug(f"[ASSEMBLER] Added text without citations ({len(clean_text)} chars)")
    
    def save_to_file(self, doc: Document, filepath: str):
        """Save document to file.
        
        Args:
            doc: Document object
            filepath: Output file path
        """
        doc.save(filepath)
        logger.info(f"[ASSEMBLER] Document saved to: {filepath}")
    
    def get_docx_bytes(self, doc: Document) -> bytes:
        """Get document as bytes for streaming.
        
        Args:
            doc: Document object
            
        Returns:
            Document bytes
        """
        from io import BytesIO
        
        logger.debug("[ASSEMBLER] Converting document to bytes")
        
        try:
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            
            docx_bytes = buffer.getvalue()
            logger.info(f"[ASSEMBLER] Document converted to {len(docx_bytes)} bytes")
            
            return docx_bytes
        except Exception as e:
            logger.error(f"[ASSEMBLER] Failed to convert document to bytes: {e}", exc_info=True)
            raise
